"""장르별 양식 자동 출력 — .docx 변환
exporter.py 분기:
- scene_dialogue: 드라마/영화/애니 — S#1. 장소 / 시간 + 캐릭터 대사
- shortdrama_episode: 숏드라마 — EP#1 + 러닝타임 + 페이월 표시
- webtoon_cuts: 웹툰 — [컷 1] 형식
- webnovel_text: 웹소설 — 줄글
- documentary_cuesheet: 다큐 — 영상/오디오 2단
- variety_cuesheet: 예능 — 코너별 큐시트
- musical_libretto: 뮤지컬 — M-1, [솔로/듀엣/합창]
- youtube_script: 유튜브 — 썸네일/첫3초/본론/콜투액션
- exhibition_plan: 전시 — 존별 평면도
- game_dialogue_data: 게임 — 엑셀 시트
- anime_scene: 애니 — 시각 묘사 강조
"""

import io
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# 한국 시나리오 표준 양식 설정 (장르 공통 베이스)
SCENARIO_FONT_BODY = "함초롬바탕"
SCENARIO_FONT_FALLBACK = "맑은 고딕"
LINE_SPACING_MULTI = 1.6


def _setup_doc_style(doc: Document, font_name: str = SCENARIO_FONT_BODY, size: int = 11):
    """문서 기본 스타일 설정 (한국 시나리오 표준)"""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size)
    style.paragraph_format.line_spacing = LINE_SPACING_MULTI


def _add_cover(doc: Document, title: str, author: str = "", episode: str = "", date: str = "", anonymous: bool = False):
    """표지 추가 — 공모전 익명 심사 룰 옵션"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\n" + title + "\n")
    run.font.size = Pt(28)
    run.bold = True

    if episode:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(episode + "\n").font.size = Pt(14)

    if author and not anonymous:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.add_run("\n\n\n작가: " + author).font.size = Pt(12)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(date or datetime.now().strftime("%Y. %m. %d.")).font.size = Pt(10)

    doc.add_page_break()


def export_to_docx(
    title: str,
    body: str,
    genre_letter: str = "B",
    export_format: str = "scene_dialogue",
    author: str = "",
    episode: str = "",
    anonymous: bool = False,
    font_name: str = SCENARIO_FONT_BODY,
) -> bytes:
    """장르 코드/양식에 맞춰 .docx 생성 → bytes 반환"""

    doc = Document()
    _setup_doc_style(doc, font_name=font_name)
    _add_cover(doc, title, author=author, episode=episode, anonymous=anonymous)

    # 양식별 분기
    if export_format == "scene_dialogue":
        _format_scene_dialogue(doc, body)
    elif export_format == "shortdrama_episode":
        _format_shortdrama(doc, body)
    elif export_format == "webtoon_cuts":
        _format_webtoon(doc, body)
    elif export_format == "webnovel_text":
        _format_webnovel(doc, body)
    elif export_format == "documentary_cuesheet":
        _format_documentary(doc, body)
    elif export_format == "variety_cuesheet":
        _format_variety(doc, body)
    elif export_format == "musical_libretto":
        _format_musical(doc, body)
    elif export_format == "youtube_script":
        _format_youtube(doc, body)
    elif export_format == "exhibition_plan":
        _format_exhibition(doc, body)
    elif export_format == "game_dialogue_data":
        _format_game(doc, body)
    elif export_format == "anime_scene":
        _format_scene_dialogue(doc, body)
    else:
        _format_scene_dialogue(doc, body)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============ 양식별 포맷터 ============

def _format_scene_dialogue(doc: Document, body: str):
    """드라마/영화/애니 — S#1. 장소 / 시간 + 캐릭터 + 대사"""
    for line in body.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # 씬 헤딩
        if stripped.startswith("S#") or stripped.startswith("씬"):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11)
            continue

        # 캐릭터 이름 (괄호 없이 짧은 줄, 다음 줄이 대사일 가능성)
        if len(stripped) < 12 and not stripped.endswith(("다", "요", "까", "지", "어", ".", "?", "!", "…")) and "(" not in stripped:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(2)
            run = p.add_run(stripped)
            run.bold = True
            continue

        # 대사 (앞뒤 따옴표 또는 들여쓰기 필요)
        if stripped.startswith(("\"", "“")) or stripped.startswith(("(", "（")):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(2.5)
            p.add_run(stripped)
            continue

        # 지문/액션
        p = doc.add_paragraph()
        p.add_run(stripped)


def _format_shortdrama(doc: Document, body: str):
    """숏드라마 — EP#1 + 러닝타임 + 페이월 표시 + 짧은 대사"""
    for line in body.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("EP#") or "[페이월" in stripped or "[CUT" in stripped:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.color.rgb = RGBColor(0x2C, 0x5B, 0xB5)
            continue
        p = doc.add_paragraph()
        p.add_run(stripped)


def _format_webtoon(doc: Document, body: str):
    """웹툰 — [컷 N] + 묘사 + 대사 + 효과음"""
    for line in body.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("[컷") or stripped.startswith("[엔딩컷") or stripped.startswith("#"):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.color.rgb = RGBColor(0x14, 0xB8, 0xA6)
            continue
        p = doc.add_paragraph()
        p.add_run(stripped)


def _format_webnovel(doc: Document, body: str):
    """웹소설 — 줄글, 단락 짧게, 챕터 헤더"""
    for paragraph in body.split("\n\n"):
        text = paragraph.strip()
        if not text:
            continue
        if text.startswith(("1화", "2화", "제1화", "제2화")) or "화. " in text[:10]:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            continue
        for line in text.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)


def _format_documentary(doc: Document, body: str):
    """다큐 — 영상/오디오 2단 큐시트 또는 구성안 줄글"""
    if "|" in body and ("영상" in body[:200] or "내레이션" in body[:200]):
        # 큐시트형 — 표
        rows = [r for r in body.split("\n") if r.strip() and "|" in r]
        if rows:
            cols = max(len(r.split("|")) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Light Grid Accent 1"
            for i, row in enumerate(rows):
                cells = [c.strip() for c in row.split("|")]
                for j, cell in enumerate(cells):
                    if j < cols:
                        table.cell(i, j).text = cell
            return

    # 구성안형 — 줄글
    for line in body.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("[챕터"):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(13)
            continue
        doc.add_paragraph(stripped)


def _format_variety(doc: Document, body: str):
    """예능 — 코너별 큐시트 + 자막 포인트"""
    for line in body.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("[코너") or stripped.startswith("[자막"):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.color.rgb = RGBColor(0xF9, 0x73, 0x16)
            continue
        if any(stripped.startswith(prefix) for prefix in ("출연:", "장소:", "구성:", "예상 자막")):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            continue
        doc.add_paragraph(stripped)


def _format_musical(doc: Document, body: str):
    """뮤지컬 — M-1. + [솔로/듀엣/합창]"""
    for line in body.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("M-") or stripped.startswith("[솔로") or stripped.startswith("[듀엣") or stripped.startswith("[합창"):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.color.rgb = RGBColor(0x8B, 0x5C, 0xF6)
            continue
        if stripped.startswith("(노래") or stripped.startswith("(VERSE") or stripped.startswith("(코러스"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(2)
            run = p.add_run(stripped)
            run.italic = True
            continue
        doc.add_paragraph(stripped)


def _format_youtube(doc: Document, body: str):
    """유튜브 — 썸네일/첫3초/후크/본론/콜투액션 구조"""
    for line in body.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if any(stripped.startswith(prefix) for prefix in ("EP", "썸네일:", "첫 3초:", "후크:", "본론:", "콜투액션:")):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            continue
        doc.add_paragraph(stripped)


def _format_exhibition(doc: Document, body: str):
    """전시 — 존별 + 동선"""
    for line in body.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("존") or stripped.startswith("Zone"):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x14, 0xB8, 0xA6)
            continue
        doc.add_paragraph(stripped)


def _format_game(doc: Document, body: str):
    """게임 — 컷씬 (대사 데이터는 별도 xlsx)"""
    if "|" in body:
        # 표 형식 (대사 데이터)
        rows = [r for r in body.split("\n") if r.strip() and "|" in r]
        if rows:
            cols = max(len(r.split("|")) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Light Grid Accent 1"
            for i, row in enumerate(rows):
                cells = [c.strip() for c in row.split("|")]
                for j, cell in enumerate(cells):
                    if j < cols:
                        table.cell(i, j).text = cell
            return

    # 컷씬 — 일반 시나리오 양식
    _format_scene_dialogue(doc, body)


def export_simple_text(content: str, filename: str = "export.txt") -> bytes:
    """텍스트 그대로 export (웹툰/웹소설 .txt 양식)"""
    return content.encode("utf-8")
