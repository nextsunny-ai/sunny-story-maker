"""첨부 파일 파싱 — .docx, .pdf, .txt, .fountain 지원"""

import io
from pathlib import Path


def parse_uploaded_file(uploaded_file) -> str:
    """Streamlit UploadedFile을 텍스트로 변환"""
    if uploaded_file is None:
        return ""

    name = uploaded_file.name.lower()
    raw = uploaded_file.read()

    if name.endswith(".docx"):
        return _parse_docx(raw)
    if name.endswith(".pdf"):
        return _parse_pdf(raw)
    if name.endswith((".txt", ".fountain", ".md", ".fdx")):
        return _parse_text(raw)
    if name.endswith(".hwp"):
        return _parse_hwp_hint()

    return _parse_text(raw)  # 마지막 시도


def _parse_docx(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return "[ERROR] python-docx 미설치. pip install python-docx"

    doc = Document(io.BytesIO(raw))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _parse_pdf(raw: bytes) -> str:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        return "[ERROR] PyPDF2 미설치. pip install PyPDF2"

    reader = PdfReader(io.BytesIO(raw))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n".join(parts)


def _parse_text(raw: bytes) -> str:
    for enc in ("utf-8", "cp949", "euc-kr", "utf-16"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _parse_hwp_hint() -> str:
    return """[HWP 파일 직접 파싱은 v1에서 미지원]

해결 방법:
1. 한글 프로그램에서 → 다른 이름으로 저장 → .docx 또는 .txt
2. 또는 https://www.hwp2docx.com 같은 무료 변환 서비스
3. 변환 후 다시 업로드해주세요.
"""


def get_word_count(text: str) -> dict:
    """글자 수 / 어절 수 / 줄 수"""
    return {
        "chars_with_spaces": len(text),
        "chars_without_spaces": len(text.replace(" ", "").replace("\n", "")),
        "words": len(text.split()),
        "lines": text.count("\n") + 1 if text else 0,
        "paragraphs": len([p for p in text.split("\n\n") if p.strip()]),
    }


def estimate_genre(text: str) -> str:
    """텍스트 패턴으로 장르 추정 (간단 규칙)"""
    lower = text.lower()
    sample = text[:3000]

    if "S#" in sample or "씬 헤딩" in sample or "/낮" in sample or "/밤" in sample:
        if "EP#" in sample or "[페이월" in sample:
            return "C"  # 숏드라마
        if "M-" in sample or "[솔로]" in sample or "[전체 합창]" in sample:
            return "I"  # 뮤지컬
        if "16부" in sample or "회차" in sample:
            return "A"  # 드라마
        return "B"  # 영화 기본

    if "[컷" in sample or "엔딩컷" in sample:
        return "F"  # 웹툰
    if "다음 화에 계속" in sample or len(text) > 4500 and "1화." in sample:
        return "H"  # 웹소설
    if "내레이션" in sample and "인터뷰" in sample:
        return "G"  # 다큐
    if "Dialog_ID" in sample or "분기조건" in sample:
        return "L"  # 게임
    if "썸네일" in sample or "첫 3초" in sample or "콜투액션" in sample:
        return "J"  # 유튜브
    if "코너" in sample and ("MC" in sample or "게스트" in sample):
        return "M"  # 예능
    if "존1" in sample or "존2" in sample or "동선" in sample:
        return "K"  # 전시
    return "B"  # 기본 = 영화


def split_into_chunks(text: str, max_chars: int = 8000) -> list[str]:
    """긴 텍스트를 청크 단위로 분할 (API 호출 제한 대응)"""
    if len(text) <= max_chars:
        return [text]

    chunks = []
    paragraphs = text.split("\n\n")
    current = ""
    for para in paragraphs:
        if len(current) + len(para) + 2 > max_chars:
            if current:
                chunks.append(current)
            current = para
        else:
            current = current + "\n\n" + para if current else para
    if current:
        chunks.append(current)
    return chunks
